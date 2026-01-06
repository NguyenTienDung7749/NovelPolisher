"""Export modules for DOCX and Markdown output."""

from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .chunking import Chunk
from .utils import log_message, ensure_dir


def export_to_markdown(
    chunks: List[Chunk],
    chunk_outputs: Dict[str, str],
    output_path: Path,
    title: str = "Polished Novel",
    mode: str = "polish_vi"
) -> Path:
    """
    Export processed chunks to Markdown file.
    
    Args:
        chunks: List of original chunks (for metadata)
        chunk_outputs: Dict mapping chunk_id to processed text
        output_path: Output file path
        title: Document title
        mode: Processing mode
        
    Returns:
        Path to created file
    """
    log_message(f"Exporting to Markdown: {output_path}")
    
    ensure_dir(output_path.parent)
    
    lines = []
    
    # Title
    mode_label = "Biên tập" if mode == "polish_vi" else "Dịch"
    lines.append(f"# {title}")
    lines.append(f"*{mode_label} bởi AI*\n")
    lines.append("---\n")
    
    current_chapter = -1
    
    for chunk in chunks:
        # Add chapter heading if new chapter
        if chunk.chapter_number != current_chapter:
            current_chapter = chunk.chapter_number
            lines.append(f"\n## Chương {chunk.chapter_number}: {chunk.chapter_title}\n")
        
        # Add part heading if multi-part
        if chunk.total_parts > 1:
            lines.append(f"\n### Phần {chunk.part_number}/{chunk.total_parts}\n")
        
        # Add processed text
        output_text = chunk_outputs.get(chunk.chunk_id, "")
        if output_text:
            lines.append(output_text)
            lines.append("")
    
    # Write file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    log_message(f"Markdown exported: {len(lines)} lines")
    return output_path


def export_to_docx(
    chunks: List[Chunk],
    chunk_outputs: Dict[str, str],
    output_path: Path,
    title: str = "Polished Novel",
    mode: str = "polish_vi"
) -> Path:
    """
    Export processed chunks to DOCX file.
    
    Args:
        chunks: List of original chunks (for metadata)
        chunk_outputs: Dict mapping chunk_id to processed text
        output_path: Output file path
        title: Document title
        mode: Processing mode
        
    Returns:
        Path to created file
    """
    log_message(f"Exporting to DOCX: {output_path}")
    
    ensure_dir(output_path.parent)
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title (Heading 1)
    mode_label = "Biên tập" if mode == "polish_vi" else "Dịch"
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f"{mode_label} bởi AI")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    current_chapter = -1
    
    for chunk in chunks:
        # Add chapter heading if new chapter (Heading 2)
        if chunk.chapter_number != current_chapter:
            current_chapter = chunk.chapter_number
            heading_text = f"Chương {chunk.chapter_number}: {chunk.chapter_title}"
            doc.add_heading(heading_text, level=1)
        
        # Add part heading if multi-part (Heading 3)
        if chunk.total_parts > 1:
            part_text = f"Phần {chunk.part_number}/{chunk.total_parts}"
            doc.add_heading(part_text, level=2)
        
        # Add processed text
        output_text = chunk_outputs.get(chunk.chunk_id, "")
        if output_text:
            # Split by paragraphs and add each
            paragraphs = output_text.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    # Handle single newlines within paragraph
                    para_text = para_text.replace('\n', ' ')
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    para.paragraph_format.space_after = Pt(6)
    
    # Save document
    doc.save(str(output_path))
    
    log_message(f"DOCX exported successfully")
    return output_path


def collect_chunk_outputs(
    chunks: List[Chunk],
    chunks_dir: Path
) -> Dict[str, str]:
    """
    Collect all chunk outputs from files.
    
    Args:
        chunks: List of chunks
        chunks_dir: Directory containing chunk output files
        
    Returns:
        Dict mapping chunk_id to output text
    """
    outputs = {}
    
    for chunk in chunks:
        chunk_file = chunks_dir / f"{chunk.chunk_id}.md"
        if chunk_file.exists():
            with open(chunk_file, 'r', encoding='utf-8') as f:
                outputs[chunk.chunk_id] = f.read()
        else:
            log_message(f"Warning: Missing output for {chunk.chunk_id}")
            outputs[chunk.chunk_id] = chunk.text  # Fallback to original
    
    return outputs
